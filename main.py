from PyQt5 import Qt

import sys
import os
import random

from PyQt5.QtWidgets import *
from PyQt5.QtGui import QPainter, QColor, QPixmap, QFont
from PyQt5 import QtCore

from pdfminer.high_level import extract_text
import docx
from fpdf import FPDF

def save_pdf_in_txt(path, name_save):
    txt = extract_text(path)

    print(txt)

    with open(f'{name_save}.txt', 'w+') as file:
        file.write(txt)
        file.close()


def save_pdf_in_word(path, name_save):
    txt = extract_text(path)

    print(txt)

    doc = docx.Document()
    par1 = doc.add_paragraph('Импортированно из pdf')

    par1.add_run(txt)

    doc.save(f'{name_save}.docx')


def save_txt_in_pdf(path, name_save):
    pdf = FPDF()
    pdf.set_font("Arial", size=12)
    pdf.add_page()
    with open(path, 'r') as file:
        for line in file:
            txt = line
            print(txt)
            pdf.cell(200, 10, txt=txt, ln=1, align="C")
        file.close()

    pdf.output(f'{name_save}.pdf')

class Example(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        u = QGridLayout(self)

        self.btn = QPushButton('Конвертировать', self)
        self.btn.clicked.connect(self.k)
        self.btn.setFont(QFont("Times", 14))
        self.btn.setStyleSheet(
            "background-color: {}".format('#00c8ff'))
        self.btn.setMaximumSize(140000, 140000)
        u.addWidget(self.btn, 0, 0)

        self.combo = QComboBox(self)
        self.combo.addItems(["txt", "pdf",
                        "docx"])
        u.addWidget(self.combo)
        
    def k(self):
        self.cwd = os.getcwd()
        fileName_choose, filetype = QFileDialog.getOpenFileName(self,  
                                    "Выбрать файл",  
                                    self.cwd, 
                                    "All Files (*);;Text Files (*.txt)")
        self.ftp = fileName_choose.split('.')[-1]
        self.f = fileName_choose
        self.tp = self.combo.currentText()
        self.nm = '.'.join((self.f.split('/')[-1]).split('.')[:-1])

        if self.ftp == 'pdf':
            if self.tp == 'txt':
                save_pdf_in_txt(self.f, self.nm)
                print('ok')
            elif self.tp == 'docx':
                save_pdf_in_word(self.f, self.nm)
                print('ok')
        if self.ftp == 'txt':
            if self.tp == 'pdf':
                save_txt_in_pdf(self.f, self.nm)
                print('ok')

if __name__ == "__main__":
    app = Qt.QApplication([])
    t = Example()
    t.show()
    app.exec_()
